"""
Adapter pour parser les documents (PDF, DOCX, TXT)
Architecture Hexagonale : Infrastructure Layer
"""

import io
import logging
from typing import List, Tuple
import fitz  # PyMuPDF
from docx import Document as DocxDocument

from ...domain.entities.document import Document, Chunk

logger = logging.getLogger(__name__)


class DocumentParserAdapter:
    """Adapter pour parser différents types de documents."""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialise l'adapter.

        Args:
            chunk_size: Taille des chunks en caractères
            chunk_overlap: Chevauchement entre chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def parse_document(self, file_bytes: bytes, filename: str) -> Document:
        """
        Parse un document et le découpe en chunks.

        Args:
            file_bytes: Contenu binaire du fichier
            filename: Nom du fichier

        Returns:
            Document avec chunks

        Raises:
            ValueError: Si le type de fichier n'est pas supporté
        """
        logger.info(f"Parsing du document: {filename}")

        # Extraire le texte selon le type
        if filename.lower().endswith(".pdf"):
            text = self._extract_text_from_pdf(file_bytes)
        elif filename.lower().endswith(".docx"):
            text = self._extract_text_from_docx(file_bytes)
        elif filename.lower().endswith(".txt"):
            text = file_bytes.decode("utf-8", errors="ignore")
        else:
            raise ValueError(
                f"Type de fichier non supporté: {filename}. "
                "Formats acceptés: PDF, DOCX, TXT"
            )

        if not text or not text.strip():
            raise ValueError(f"Le document {filename} ne contient pas de texte")

        # Découper en chunks
        chunks = self._create_chunks(text, filename)

        # Créer l'entité Document
        document = Document(
            filename=filename,
            content=text,
            chunks=chunks,
            metadata={
                "file_type": self._get_file_type(filename),
                "text_length": len(text),
                "chunks_count": len(chunks)
            }
        )

        logger.info(
            f"Document {filename} parsé: {len(text)} caractères, "
            f"{len(chunks)} chunks"
        )

        return document

    def _extract_text_from_pdf(self, file_bytes: bytes) -> str:
        """Extrait le texte d'un fichier PDF."""
        text = ""
        try:
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                for page in doc:
                    text += page.get_text()
        except Exception as e:
            logger.error(f"Erreur extraction PDF: {e}")
            raise ValueError(f"Erreur lors de l'extraction du PDF: {e}")
        return text

    def _extract_text_from_docx(self, file_bytes: bytes) -> str:
        """Extrait le texte d'un fichier DOCX, y compris les tableaux."""
        text_parts = []
        try:
            doc = DocxDocument(io.BytesIO(file_bytes))

            # Extraire les paragraphes
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extraire le contenu des tableaux
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    row_text = " | ".join(row_cells)
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    text_parts.append("\n".join(table_text))

        except Exception as e:
            logger.error(f"Erreur extraction DOCX: {e}")
            raise ValueError(f"Erreur lors de l'extraction du DOCX: {e}")

        return "\n".join(text_parts)

    def _create_chunks(self, text: str, filename: str) -> List[Chunk]:
        """
        Découpe le texte en chunks avec chevauchement.

        Args:
            text: Texte complet
            filename: Nom du fichier source

        Returns:
            Liste de chunks
        """
        chunks = []
        start = 0

        while start < len(text):
            end = start + self.chunk_size
            chunk_text = text[start:end]

            # Essayer de couper à la fin d'une phrase
            if end < len(text):
                last_period = chunk_text.rfind(".")
                last_newline = chunk_text.rfind("\n")
                break_point = max(last_period, last_newline)

                if break_point > self.chunk_size // 2:
                    chunk_text = chunk_text[:break_point + 1]
                    end = start + break_point + 1

            # Créer le chunk
            chunk = Chunk(
                text=chunk_text.strip(),
                metadata={
                    "filename": filename,
                    "chunk_index": len(chunks),
                    "start_char": start,
                    "end_char": end
                }
            )
            chunks.append(chunk)

            # Avancer avec chevauchement
            start = end - self.chunk_overlap if end < len(text) else end

        return chunks

    def _get_file_type(self, filename: str) -> str:
        """Retourne le type de fichier."""
        if filename.lower().endswith(".pdf"):
            return "pdf"
        elif filename.lower().endswith(".docx"):
            return "docx"
        elif filename.lower().endswith(".txt"):
            return "txt"
        return "unknown"
